from docx import Document

doc = Document()

doc.add_heading('Методичні вказівки та лабораторна робота', 0)
doc.add_heading('Тема: Оцінка техніко-економічних показників розробки ПЗ', level=1)
doc.add_paragraph('Дисципліна: Стандартизація та сертифікація програмного забезпечення')

doc.add_heading('1. Мета роботи', level=2)
doc.add_paragraph(
"Формування навичок оцінювання техніко-економічних показників програмного забезпечення та застосування моделей COCOMO."
)

doc.add_heading('2. Основні формули', level=2)
doc.add_paragraph("T = n * t")
doc.add_paragraph("ФОП = n * ЗП * t")
doc.add_paragraph("В = ФОП + ДВ")
doc.add_paragraph("Tок = В / Е")
doc.add_paragraph("R = (Е / В) * 100%")

doc.add_heading('3. Додаткове завдання (COCOMO)', level=2)
doc.add_paragraph(
"Виконати розрахунки за базовою та проміжною моделями COCOMO: C, T, N, P."
)

doc.add_heading('4. Формули COCOMO', level=3)
doc.add_paragraph("C = a * (KLOC)^b")
doc.add_paragraph("T = c * (C)^d")
doc.add_paragraph("C = a * (KLOC)^b * EAF")
doc.add_paragraph(f"N = C / T")
doc.add_paragraph("P = KLOC / C")

doc.add_heading('5. Коефіцієнти моделі COCOMO', level=2)
table = doc.add_table(rows=1, cols=5)
hdr = table.rows[0].cells
hdr[0].text = 'Тип'
hdr[1].text = 'a'
hdr[2].text = 'b'
hdr[3].text = 'c'
hdr[4].text = 'd'

rows = [
("Органічний","2.4","1.05","2.5","0.38"),
("Напівзалежний","3.0","1.12","2.5","0.35"),
("Вбудований","3.6","1.20","2.5","0.32"),
]

for r in rows:
    row_cells = table.add_row().cells
    for i in range(5):
        row_cells[i].text = r[i]

doc.add_heading('6. Таблиця EAF (коефіцієнти)', level=2)
doc.add_paragraph("Приклад факторів множника складності (EAF):")

table2 = doc.add_table(rows=1, cols=2)
h = table2.rows[0].cells
h[0].text = 'Фактор'
h[1].text = 'Значення'

eaf_rows = [
("Досвід команди", "0.8 – 1.3"),
("Складність ПЗ", "0.9 – 1.5"),
("Інструменти розробки", "0.8 – 1.2"),
("Надійність вимог", "0.9 – 1.4"),
]

for r in eaf_rows:
    row = table2.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]

doc.add_heading('7. Шаблон для студентів', level=2)
doc.add_paragraph("Варіант: ________")
doc.add_paragraph("KLOC: ________")
doc.add_paragraph("Тип проєкту: ________")
doc.add_paragraph("EAF: ________")

doc.add_heading('Розрахунки:', level=3)
doc.add_paragraph("C (базова модель): ____________________")
doc.add_paragraph("C (проміжна модель): __________________")
doc.add_paragraph("T: _________________________________")
doc.add_paragraph("N: _________________________________")
doc.add_paragraph("P: _________________________________")

doc.add_heading('Висновок:', level=3)
doc.add_paragraph("______________________________________________")
doc.add_paragraph("______________________________________________")

file_path = "ЛР_COCOMO_ТЕП_ПЗ.docx"
doc.save(file_path)

